import pandas as pd
from docx import Document
from typing import List, Tuple, Dict
import os
import datetime
import time
from tqdm import tqdm


def get_data(data_path: str) -> Tuple[int, pd.DataFrame, List[str]]:
    df = pd.read_excel(data_path, engine='openpyxl')
    return len(df), df, df.columns.tolist()


def fill_invitations(template_path: str, data, key_format: str) -> Document:
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if " " not in str(key_format):
                key = key.replace(' ', '_')

            key = key_format.replace(key_format[1:len(key_format) - 1], key)
            opening_bracket = key_format[0]
            closing_bracket = key_format[-1]

            if key in paragraph.text:
                temp = ""
                for run in paragraph.runs:
                    if run.text == opening_bracket or temp != "":
                        temp += run.text
                        run.clear()

                    if run.text.count(opening_bracket) > run.text.count(closing_bracket):
                        temp_2 = run.text.split(opening_bracket)
                        run.text = temp_2[0]
                        temp += (opening_bracket + temp_2[1])

                    if closing_bracket in temp and temp.count(opening_bracket) == temp.count(closing_bracket):
                        run.text = temp + run.text
                        if temp.count(opening_bracket) > 1:
                            for key_, value_ in data.items():
                                run.text = run.text.replace(f"{opening_bracket}{key_}{closing_bracket}", str(value_))
                        else:
                            run.text = run.text.replace(f"{key}", str(value))
                        temp = ""

                    if key in run.text:
                        run.text = run.text.replace(f"{key}", str(value))


    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:  # Iterate through paragraphs inside the cell
                    # Iterate through runs inside the paragraph
                    for key, value in data.items():
                        if " " not in str(key_format):
                            key = key.replace(' ', '_')

                        key = key_format.replace(key_format[1:len(key_format) - 1], key)
                        opening_bracket = key_format[0]
                        closing_bracket = key_format[-1]
                        
                        if key in paragraph.text:
                            temp = ""
                            for run in paragraph.runs:
                                if run.text == opening_bracket or temp != "":
                                    temp += run.text
                                    run.clear()

                                if run.text.count(opening_bracket) > run.text.count(closing_bracket):
                                    temp_2 = run.text.split(opening_bracket)
                                    run.text = temp_2[0]
                                    temp += (opening_bracket + temp_2[1])

                                if closing_bracket in temp and temp.count(opening_bracket) == temp.count(
                                        closing_bracket):
                                    run.text = temp + run.text
                                    if temp.count(opening_bracket) > 1:
                                        for key_, value_ in data.items():
                                            run.text = run.text.replace(f"{opening_bracket}{key_}{closing_bracket}", str(value_))
                                    else:
                                        run.text = run.text.replace(f"{key}", str(value))
                                    temp = ""

                                if key in run.text:
                                    run.text = run.text.replace(f"{key}", str(value))
    return doc


def merge_documents(documents: List[Document]) -> Document:
    merged_doc = Document()
    for i, doc in enumerate(documents):
        # Append each element of the current document to the merged document
        if i != len(documents) - 1:
            doc.add_page_break()

        for element in doc.element.body:
            merged_doc.element.body.append(element)

    return merged_doc



def main():
    pass

if __name__ == "__main__":
    main()
