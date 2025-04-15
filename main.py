import pandas as pd
from docx import Document
from typing import List, Tuple, Dict
import os
import datetime
import time
from tqdm import tqdm


def get_data(data_path: str) -> Tuple[int, pd.DataFrame, List[str]]:
    df = pd.read_excel(data_path, engine='openpyxl')
    return len(df), df, df.columns.tolist()


def fill_invitations(template_path: str, data, debug:bool = False) -> Document:
    doc = Document(template_path)
    stream = False
    table_stream = False
    for paragraph in doc.paragraphs:
        if debug:
            print(paragraph.text)
            print("-" * 10)

        for key, value in data.items():
            key = key.replace(' ', '_')
            key = f"«{key}»"

            if key == f"«Địa_chỉ_thửa_đất_mới»":
                if "Xã" in value:
                    value = value.replace("Xã", "xã")

            if key in paragraph.text:
                if debug:
                    print(f"Found {key}")
                    print(f"Replacing: {key} with {str(value)}")

                # Clear runs and re-add the modified text
                temp = ""

                for run in paragraph.runs:
                    if debug or stream:
                        print("First:", run.text)

                    if run.text == "«" or temp != "":
                        if debug or stream:
                            print("Temp open:", temp)
                        temp += run.text
                        run.clear()

                    if "«" in run.text and "»" not in run.text:
                        if debug or stream:
                            print("Temp open 2:", temp)
                        temp_2 = run.text.split("«")
                        if debug or stream:
                            print("Split list:", temp_2)
                        run.text = temp_2[0]
                        if debug or stream:
                            print("Curr run: ", run.text)
                        temp += ("«" + temp_2[1])

                    if "«" in run.text and "»" in run.text:
                        run.text = run.text.replace(f"{key}", str(value))

                    if "»" in temp:
                        if debug or stream:
                            print("Temp closing:", temp)
                        run.text = temp + run.text
                        if debug or stream:
                            print("Addition:", run.text)
                        run.text = run.text.replace(f"{key}", str(value))
                        if debug or stream:
                            print("End:", run.text)
                        temp = ""
                    if debug or stream:
                        print()
                        print(paragraph.text)
                        print("-" * 50)
                # for run in paragraph.runs:
                #     run.text = run.text.replace(f"{key}", str(value))

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:  # Iterate through paragraphs inside the cell
                    # Iterate through runs inside the paragraph
                    for key, value in data.items():
                        key = key.replace(' ', '_')
                        key = f"«{key}»"

                        if key == f"«Địa_chỉ_thửa_đất_mới»":
                            if "Xã" in value:
                                value = value.replace("Xã", "xã")

                        temp = ""
                        for run in paragraph.runs:
                            if debug or table_stream:
                                print(f"Run text:", run.text)

                            if run.text == "«" or temp != "":
                                if debug or stream:
                                    print("Temp open:", temp)
                                temp += run.text
                                run.clear()

                            if "«" in run.text and "»" not in run.text:
                                if debug or stream:
                                    print("Temp open 2:", temp)
                                temp_2 = run.text.split("«")
                                if debug or stream:
                                    print("Split list:", temp_2)
                                run.text = temp_2[0]
                                if debug or stream:
                                    print("Curr run: ", run.text)
                                temp += ("«" + temp_2[1])

                            if "«" in run.text and "»" in run.text:
                                run.text = run.text.replace(f"{key}", str(value))

                            if "»" in temp:
                                if debug or stream:
                                    print("Temp closing:", temp)
                                run.text = temp + run.text
                                if debug or stream:
                                    print("Addition:", run.text)
                                run.text = run.text.replace(f"{key}", str(value))
                                if debug or stream:
                                    print("End:", run.text)
                                temp = ""
    return doc


def run(data_path: str,
        template_path: str,
        output_folder: str,
        data: Dict[str, str],
        isMerged: bool = False,
        test_iterations: int = 0,
        debug: bool = False) -> None:

    iterations, df, columns = get_data(data_path)

    if test_iterations >= 1:
        iterations = test_iterations

    documents = []

    for iteration in tqdm(range(iterations), desc="Document"):
        data = {}
        row = df.iloc[iteration, :]
        for column in columns:
            data[column.strip()] = row.loc[column]

        if debug:
            print(data)

        # Fill the invitation and get the document
        doc = fill_invitations(template_path, data, debug)

        if not isMerged:
            file_path = os.path.join(output_folder, f"THÔNG BÁO {row.loc['STT']}.docx")
            doc.save(file_path)
            if debug:
                print(f"Document saved as: {file_path}")
        else:
            documents.append(doc)


def main():
    data_path = r"info.xlsx"
    template_path = r"format1.docx"
    output_folder = r"results/"

    data = {

    }

    run(
        data_path=data_path,
        template_path=template_path,
        output_folder=output_folder,
        data=data,
        isMerged=False,
        test_iterations=0,
        debug=False
    )


if __name__ == "__main__":
    start = time.time()
    main()
    print(f"Elapsed time: {round(time.time() - start, 2)}s")
