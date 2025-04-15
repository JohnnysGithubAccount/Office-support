import pandas as pd
from docx import Document
from typing import List, Tuple, Dict
import os
import time
from tqdm import tqdm


def get_data(data_path: str) -> Tuple[int, pd.DataFrame, List[str]]:
    df = pd.read_excel(data_path, engine='openpyxl')
    return len(df), df, df.columns.tolist()


def fill_invitations(template_path: str, data, debug:bool = False) -> Document:
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        combined_text = ''.join(run.text for run in paragraph.runs)

        for key, value in data.items():
            key = f"[{key}]"
            if key in paragraph.text:
                combined_text = combined_text.replace(key, str(value))

                for run in paragraph.runs:
                    run.clear()

                for part in combined_text.split('\n'):
                    new_run = paragraph.add_run(part)
                    new_run.bold = False
    return doc


def merge_documents(documents: List[Document]) -> Document:
    merged_doc = Document()
    for i, doc in enumerate(documents):
        if i != len(documents) - 1:
            doc.add_page_break()

        for element in doc.element.body:
            merged_doc.element.body.append(element)

    return merged_doc


def run(data_path: str,
        template_path: str,
        output_folder: str,
        data: Dict[str, str],
        isMerged: bool = False,
        test_iterations: int = 0,
        debug: bool = False) -> None:

    iterations, df, columns = get_data(data_path)

    if test_iterations >= 1:
        iterations = test_iterations

    documents = []

    for iteration in tqdm(range(iterations), desc="Document"):
        row = df.iloc[iteration, :]
        for column in columns:
            data[column.strip()] = row.loc[column]

        doc = fill_invitations(template_path, data, debug)

        if not isMerged:
            file_path = os.path.join(output_folder, f"THÔNG BÁO {row.loc['STT']}.docx")
            doc.save(file_path)
        else:
            documents.append(doc)

    if len(documents) > 1 and isMerged:
        merged_doc = merge_documents(documents)

        merged_output_path = os.path.join(output_folder, "merged_invitations.docx")
        merged_doc.save(merged_output_path)


def main():
    pass


if __name__ == "__main__":
    start = time.time()
    main()
    print(f"Elapsed time: {round(time.time() - start, 2)}s")
